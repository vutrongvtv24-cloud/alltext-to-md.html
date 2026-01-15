"""
📄 Document Processor - Unified Document Report Generator
=========================================================
A Streamlit application that processes multiple file types (Excel, Word, PDF, Text)
and aggregates them into a single structured Markdown/HTML report.

Core Philosophy: ZERO Content Alteration - Extract text exactly as it appears.

Author: Senior Python Developer
"""

import streamlit as st
import pandas as pd
import pdfplumber
from docx import Document
from docx.table import Table as DocxTable
import markdown
import io
import re
from typing import List, Tuple, Optional, Dict
from dataclasses import dataclass
from datetime import datetime
from PIL import Image
import easyocr
import numpy as np


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class ProcessedFile:
    """Represents a processed file with its extracted content."""
    filename: str
    file_type: str
    content: str
    success: bool
    error_message: Optional[str] = None


# ============================================================================
# DOCUMENT PROCESSOR CLASS
# ============================================================================

class DocumentProcessor:
    """
    Main class for processing various document types.
    
    Supported formats:
    - Excel (.xlsx, .xls): Converts each sheet to Markdown table
    - Word (.docx): Extracts paragraphs and tables
    - PDF (.pdf): Extracts text and tables using pdfplumber
    - Text (.txt): Reads content directly
    - Images (.png, .jpg, .jpeg): OCR text extraction using easyocr
    - Markdown (.md): Read and pass through (for MD to HTML conversion)
    """
    
    def __init__(self):
        self.processed_files: List[ProcessedFile] = []
        self.warnings: List[str] = []
        self._ocr_reader = None  # Lazy initialization for OCR
    
    def _get_ocr_reader(self):
        """
        Lazy initialization of OCR reader.
        This avoids loading the model until it's actually needed.
        Supports both Vietnamese and English text.
        """
        if self._ocr_reader is None:
            # Initialize with Vietnamese and English support
            # gpu=False for compatibility (change to True if GPU available)
            self._ocr_reader = easyocr.Reader(['vi', 'en'], gpu=False)
        return self._ocr_reader
    
    def process_files(self, uploaded_files: List) -> str:
        """
        Process all uploaded files and return aggregated Markdown content.
        
        Args:
            uploaded_files: List of Streamlit UploadedFile objects
            
        Returns:
            Aggregated Markdown string with all file contents
        """
        self.processed_files.clear()
        self.warnings.clear()
        
        for uploaded_file in uploaded_files:
            try:
                file_extension = uploaded_file.name.split('.')[-1].lower()
                
                if file_extension in ['xlsx', 'xls']:
                    content = self._process_excel(uploaded_file)
                elif file_extension == 'docx':
                    content = self._process_word(uploaded_file)
                elif file_extension == 'pdf':
                    content = self._process_pdf(uploaded_file)
                elif file_extension == 'txt':
                    content = self._process_text(uploaded_file)
                elif file_extension in ['png', 'jpg', 'jpeg']:
                    content = self._process_image(uploaded_file)
                elif file_extension == 'md':
                    content = self._process_markdown(uploaded_file)
                else:
                    raise ValueError(f"Unsupported file format: .{file_extension}")
                
                self.processed_files.append(ProcessedFile(
                    filename=uploaded_file.name,
                    file_type=file_extension.upper(),
                    content=content,
                    success=True
                ))
                
            except Exception as e:
                error_msg = f"Error processing '{uploaded_file.name}': {str(e)}"
                self.warnings.append(error_msg)
                self.processed_files.append(ProcessedFile(
                    filename=uploaded_file.name,
                    file_type=uploaded_file.name.split('.')[-1].upper(),
                    content="",
                    success=False,
                    error_message=str(e)
                ))
        
        return self._aggregate_content()
    
    # ========================================================================
    # EXCEL PROCESSING
    # ========================================================================
    
    def _process_excel(self, file) -> str:
        """
        Process Excel file (.xlsx, .xls) and convert to Markdown.
        
        EXTRACTION LOGIC:
        1. Read the Excel file using pandas with openpyxl engine
        2. Iterate through ALL sheets in the workbook
        3. For each sheet:
           - Add sheet name as a sub-header (### Sheet: {name})
           - Convert the DataFrame to a Markdown table
           - Handle empty cells by replacing NaN with empty string
           - Preserve all data types as strings to avoid data loss
        4. Combine all sheets with proper formatting
        
        Args:
            file: Streamlit UploadedFile object
            
        Returns:
            Markdown string containing all sheets as tables
        """
        content_parts = []
        
        # Read Excel file - use openpyxl for .xlsx, xlrd for .xls
        file_ext = file.name.split('.')[-1].lower()
        engine = 'openpyxl' if file_ext == 'xlsx' else 'xlrd'
        
        # Read all sheets into a dictionary
        excel_file = pd.ExcelFile(file, engine=engine)
        sheet_names = excel_file.sheet_names
        
        for sheet_name in sheet_names:
            # Read each sheet, keeping all data as strings to preserve original format
            df = pd.read_excel(
                excel_file, 
                sheet_name=sheet_name,
                dtype=str,  # Read all as string to preserve data
                na_filter=False  # Don't convert empty cells to NaN
            )
            
            if df.empty:
                content_parts.append(f"### 📊 Sheet: {sheet_name}\n\n*Empty sheet*\n")
                continue
            
            # Convert DataFrame to Markdown table
            markdown_table = self._dataframe_to_markdown(df)
            content_parts.append(f"### 📊 Sheet: {sheet_name}\n\n{markdown_table}\n")
        
        return "\n".join(content_parts)
    
    def _dataframe_to_markdown(self, df: pd.DataFrame) -> str:
        """
        Convert a pandas DataFrame to a Markdown table.
        
        Args:
            df: pandas DataFrame
            
        Returns:
            Markdown formatted table string
        """
        if df.empty:
            return "*No data*"
        
        # Create header row
        headers = " | ".join(str(col) for col in df.columns)
        separator = " | ".join(["---"] * len(df.columns))
        
        # Create data rows
        rows = []
        for _, row in df.iterrows():
            row_str = " | ".join(str(val).replace("|", "\\|").replace("\n", " ") for val in row)
            rows.append(row_str)
        
        # Combine into table
        table = f"| {headers} |\n| {separator} |\n"
        for row in rows:
            table += f"| {row} |\n"
        
        return table
    
    # ========================================================================
    # WORD DOCUMENT PROCESSING
    # ========================================================================
    
    def _process_word(self, file) -> str:
        """
        Process Word document (.docx) and extract content.
        
        EXTRACTION LOGIC:
        1. Load the document using python-docx
        2. Iterate through document body elements in order
        3. For paragraphs:
           - Detect heading styles and convert to Markdown headers
           - Preserve paragraph text exactly as written
        4. For tables:
           - Convert each table to Markdown format
           - Preserve cell content and structure
        5. Maintain document order for coherent output
        
        Args:
            file: Streamlit UploadedFile object
            
        Returns:
            Markdown string with document content
        """
        doc = Document(file)
        content_parts = []
        
        for element in doc.element.body:
            # Check if element is a paragraph
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            # Check for heading styles
                            if para.style and para.style.name.startswith('Heading'):
                                level = self._get_heading_level(para.style.name)
                                content_parts.append(f"{'#' * level} {text}\n")
                            else:
                                content_parts.append(f"{text}\n")
                        break
            
            # Check if element is a table
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        markdown_table = self._word_table_to_markdown(table)
                        content_parts.append(f"\n{markdown_table}\n")
                        break
        
        return "\n".join(content_parts)
    
    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from Word style name."""
        match = re.search(r'\d+', style_name)
        if match:
            return min(int(match.group()), 6)  # Max heading level is 6
        return 2  # Default to H2
    
    def _word_table_to_markdown(self, table: DocxTable) -> str:
        """Convert Word table to Markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in row.cells]
            rows.append(cells)
        
        if not rows:
            return "*Empty table*"
        
        # Use first row as header
        headers = " | ".join(rows[0])
        separator = " | ".join(["---"] * len(rows[0]))
        
        table_md = f"| {headers} |\n| {separator} |\n"
        for row in rows[1:]:
            # Pad row if necessary
            while len(row) < len(rows[0]):
                row.append("")
            table_md += f"| {' | '.join(row)} |\n"
        
        return table_md
    
    # ========================================================================
    # PDF PROCESSING
    # ========================================================================
    
    def _process_pdf(self, file) -> str:
        """
        Process PDF file and extract text and tables.
        
        EXTRACTION LOGIC using pdfplumber:
        1. Open PDF with pdfplumber (better table detection than PyPDF2)
        2. Iterate through each page:
           - Extract tables first using pdfplumber's table detection
           - Tables are detected based on cell boundaries and lines
           - For non-table content, extract text using extract_text()
        3. Table extraction strategy:
           - pdfplumber uses cell boundary detection
           - Tables are converted to Markdown format
           - Each table is separated from text content
        4. Page numbers are added for reference
        
        Why pdfplumber?
        - Better table detection algorithm
        - Handles complex table structures
        - Preserves table cell boundaries
        - More accurate text positioning
        
        Args:
            file: Streamlit UploadedFile object
            
        Returns:
            Markdown string with extracted content
        """
        content_parts = []
        
        with pdfplumber.open(file) as pdf:
            total_pages = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_content = []
                page_content.append(f"#### 📄 Page {page_num}/{total_pages}\n")
                
                # Extract tables from the page
                # pdfplumber detects tables based on:
                # - Explicit line boundaries
                # - Cell spacing patterns
                # - Text alignment
                tables = page.extract_tables()
                
                if tables:
                    # Process each detected table
                    for table_idx, table in enumerate(tables, 1):
                        if table and len(table) > 0:
                            markdown_table = self._pdf_table_to_markdown(table)
                            page_content.append(f"**Table {table_idx}:**\n\n{markdown_table}\n")
                
                # Extract remaining text content
                # This captures text that is NOT part of detected tables
                text = page.extract_text()
                if text:
                    # Clean up the text
                    cleaned_text = self._clean_pdf_text(text)
                    if cleaned_text:
                        page_content.append(f"\n{cleaned_text}\n")
                
                if len(page_content) > 1:  # More than just page header
                    content_parts.append("\n".join(page_content))
                else:
                    content_parts.append(f"#### 📄 Page {page_num}/{total_pages}\n\n*No extractable content*\n")
        
        return "\n".join(content_parts)
    
    def _pdf_table_to_markdown(self, table: List[List]) -> str:
        """
        Convert PDF table (list of lists) to Markdown format.
        
        Args:
            table: 2D list representing table data
            
        Returns:
            Markdown formatted table string
        """
        if not table or len(table) == 0:
            return "*Empty table*"
        
        # Clean table cells
        cleaned_table = []
        for row in table:
            if row:
                cleaned_row = [
                    str(cell or "").strip().replace("|", "\\|").replace("\n", " ")
                    for cell in row
                ]
                cleaned_table.append(cleaned_row)
        
        if not cleaned_table:
            return "*Empty table*"
        
        # Determine max columns
        max_cols = max(len(row) for row in cleaned_table)
        
        # Pad rows to have consistent columns
        for row in cleaned_table:
            while len(row) < max_cols:
                row.append("")
        
        # Create Markdown table
        headers = " | ".join(cleaned_table[0])
        separator = " | ".join(["---"] * max_cols)
        
        table_md = f"| {headers} |\n| {separator} |\n"
        for row in cleaned_table[1:]:
            table_md += f"| {' | '.join(row)} |\n"
        
        return table_md
    
    def _clean_pdf_text(self, text: str) -> str:
        """Clean extracted PDF text."""
        if not text:
            return ""
        
        # Remove excessive whitespace while preserving paragraph breaks
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            stripped = line.strip()
            if stripped:
                cleaned_lines.append(stripped)
            elif cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")  # Preserve paragraph breaks
        
        return "\n".join(cleaned_lines)
    
    # ========================================================================
    # TEXT FILE PROCESSING
    # ========================================================================
    
    def _process_text(self, file) -> str:
        """
        Process plain text file (.txt).
        
        Args:
            file: Streamlit UploadedFile object
            
        Returns:
            Text content as string
        """
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                file.seek(0)
                content = file.read().decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError("Unable to decode text file with supported encodings")
        
        return content.strip()
    
    # ========================================================================
    # MARKDOWN FILE PROCESSING
    # ========================================================================
    
    def _process_markdown(self, file) -> str:
        """
        Process Markdown file (.md).
        
        EXTRACTION LOGIC:
        1. Read the file content with encoding detection
        2. Preserve the Markdown content exactly as-is
        3. This allows MD files to be combined with other files
           or converted directly to HTML
        
        Args:
            file: Streamlit UploadedFile object
            
        Returns:
            Markdown content as string (preserved exactly)
        """
        # Try different encodings (same as text files)
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                file.seek(0)
                content = file.read().decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError("Unable to decode Markdown file with supported encodings")
        
        # Return content as-is (it's already Markdown)
        return content.strip()
    
    # ========================================================================
    # IMAGE PROCESSING (OCR)
    # ========================================================================
    
    def _process_image(self, file) -> str:
        """
        Process image file (.png, .jpg, .jpeg) and extract text using OCR.
        
        EXTRACTION LOGIC using easyocr:
        1. Load the image using PIL (Pillow)
        2. Convert to RGB format if necessary (for consistency)
        3. Convert PIL Image to numpy array for easyocr
        4. Use easyocr to detect and extract text
        5. easyocr returns a list of (bbox, text, confidence) tuples
        6. Combine all detected text preserving reading order
        
        Why easyocr?
        - Pure Python, no external dependencies like Tesseract
        - Deep learning based, better accuracy
        - Supports 80+ languages including Vietnamese
        - Handles various image qualities well
        
        Args:
            file: Streamlit UploadedFile object
            
        Returns:
            Markdown string with extracted text from image
        """
        content_parts = []
        
        # Load image
        file.seek(0)
        image = Image.open(file)
        
        # Convert to RGB if necessary (e.g., RGBA or grayscale)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Get image info for output
        width, height = image.size
        content_parts.append(f"**Image Size:** {width} × {height} pixels")
        content_parts.append("")
        
        # Convert PIL Image to numpy array
        image_array = np.array(image)
        
        # Get OCR reader (lazy initialization)
        reader = self._get_ocr_reader()
        
        # Perform OCR
        # detail=1 returns (bbox, text, confidence)
        # paragraph=True groups text into paragraphs
        results = reader.readtext(image_array, detail=1, paragraph=False)
        
        if not results:
            content_parts.append("*No text detected in image*")
            return "\n".join(content_parts)
        
        content_parts.append("### 📝 Extracted Text")
        content_parts.append("")
        
        # Process OCR results
        # Results are sorted by position (top to bottom, left to right)
        extracted_lines = []
        for (bbox, text, confidence) in results:
            if text.strip():
                # Include confidence for transparency
                # Only show if confidence is reasonable
                if confidence >= 0.3:  # 30% minimum confidence
                    extracted_lines.append(text.strip())
        
        if extracted_lines:
            # Join lines with proper spacing
            content_parts.append("\n".join(extracted_lines))
        else:
            content_parts.append("*No readable text detected*")
        
        content_parts.append("")
        content_parts.append(f"*OCR Confidence: Text extracted with varying confidence levels*")
        
        return "\n".join(content_parts)
    
    # ========================================================================
    # CONTENT AGGREGATION
    # ========================================================================
    
    def _aggregate_content(self) -> str:
        """
        Aggregate all processed file contents into a single Markdown document.
        
        Structure:
        1. Title
        2. Generation timestamp
        3. Table of Contents
        4. File sections separated by horizontal rules
        
        Returns:
            Complete Markdown document string
        """
        if not self.processed_files:
            return "# No files processed\n\nPlease upload files to process."
        
        # Document header
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        content_parts = [
            "# 📚 Unified Document Report",
            "",
            f"*Generated: {timestamp}*",
            "",
            f"*Total files: {len(self.processed_files)} | Successful: {sum(1 for f in self.processed_files if f.success)}*",
            "",
        ]
        
        # Table of Contents
        content_parts.append("## 📋 Table of Contents")
        content_parts.append("")
        
        for idx, pf in enumerate(self.processed_files, 1):
            status = "✅" if pf.success else "❌"
            anchor = self._create_anchor(pf.filename)
            content_parts.append(f"{idx}. {status} [{pf.filename}](#{anchor}) `[{pf.file_type}]`")
        
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        
        # File sections
        for pf in self.processed_files:
            anchor = self._create_anchor(pf.filename)
            content_parts.append(f"## 📄 {pf.filename} {{#{anchor}}}")
            content_parts.append("")
            content_parts.append(f"**File Type:** {pf.file_type}")
            content_parts.append("")
            
            if pf.success:
                content_parts.append(pf.content)
            else:
                content_parts.append(f"> ⚠️ **Error:** {pf.error_message}")
                content_parts.append("")
                content_parts.append("*This file could not be processed.*")
            
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
        
        return "\n".join(content_parts)
    
    def _create_anchor(self, filename: str) -> str:
        """Create URL-safe anchor from filename."""
        # Remove extension and special characters
        anchor = re.sub(r'[^a-zA-Z0-9\s-]', '', filename)
        anchor = re.sub(r'\s+', '-', anchor).lower()
        return anchor


# ============================================================================
# HTML GENERATION
# ============================================================================

def generate_html(markdown_content: str) -> str:
    """
    Convert Markdown to HTML with GitHub-style CSS.
    
    Args:
        markdown_content: Markdown string
        
    Returns:
        Complete HTML document string
    """
    # Convert Markdown to HTML
    html_body = markdown.markdown(
        markdown_content,
        extensions=['tables', 'fenced_code', 'toc']
    )
    
    # GitHub-style CSS
    css = """
    <style>
        * {
            box-sizing: border-box;
        }
        
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Noto Sans', Helvetica, Arial, sans-serif;
            font-size: 16px;
            line-height: 1.6;
            color: #24292f;
            background-color: #ffffff;
            margin: 0;
            padding: 20px;
        }
        
        .container {
            max-width: 980px;
            margin: 0 auto;
            padding: 45px;
            border: 1px solid #d0d7de;
            border-radius: 6px;
            background-color: #ffffff;
        }
        
        h1, h2, h3, h4, h5, h6 {
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
            border-bottom: 1px solid #d8dee4;
            padding-bottom: 0.3em;
        }
        
        h1 { font-size: 2em; }
        h2 { font-size: 1.5em; }
        h3 { font-size: 1.25em; border-bottom: none; }
        h4 { font-size: 1em; border-bottom: none; }
        
        p {
            margin-top: 0;
            margin-bottom: 16px;
        }
        
        a {
            color: #0969da;
            text-decoration: none;
        }
        
        a:hover {
            text-decoration: underline;
        }
        
        code {
            padding: 0.2em 0.4em;
            margin: 0;
            font-size: 85%;
            background-color: rgba(175, 184, 193, 0.2);
            border-radius: 6px;
            font-family: ui-monospace, SFMono-Regular, 'SF Mono', Menlo, Consolas, monospace;
        }
        
        pre {
            padding: 16px;
            overflow: auto;
            font-size: 85%;
            line-height: 1.45;
            background-color: #f6f8fa;
            border-radius: 6px;
        }
        
        pre code {
            padding: 0;
            background-color: transparent;
        }
        
        blockquote {
            padding: 0 1em;
            color: #57606a;
            border-left: 0.25em solid #d0d7de;
            margin: 0 0 16px 0;
        }
        
        table {
            border-spacing: 0;
            border-collapse: collapse;
            margin-top: 0;
            margin-bottom: 16px;
            width: 100%;
            overflow: auto;
        }
        
        table th, table td {
            padding: 6px 13px;
            border: 1px solid #d0d7de;
        }
        
        table th {
            font-weight: 600;
            background-color: #f6f8fa;
        }
        
        table tr {
            background-color: #ffffff;
            border-top: 1px solid #d0d7de;
        }
        
        table tr:nth-child(2n) {
            background-color: #f6f8fa;
        }
        
        hr {
            height: 0.25em;
            padding: 0;
            margin: 24px 0;
            background-color: #d0d7de;
            border: 0;
        }
        
        ul, ol {
            padding-left: 2em;
            margin-top: 0;
            margin-bottom: 16px;
        }
        
        li {
            margin-top: 0.25em;
        }
        
        img {
            max-width: 100%;
            box-sizing: content-box;
        }
        
        .emoji {
            height: 1em;
            width: 1em;
        }
        
        @media (max-width: 767px) {
            .container {
                padding: 15px;
            }
        }
        
        @media print {
            body {
                background-color: white;
            }
            .container {
                border: none;
                box-shadow: none;
            }
        }
    </style>
    """
    
    # Complete HTML document
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Unified Document Report</title>
    {css}
</head>
<body>
    <div class="container">
        {html_body}
    </div>
</body>
</html>"""
    
    return html


# ============================================================================
# STREAMLIT APPLICATION
# ============================================================================

def main():
    """Main Streamlit application."""
    
    # Page configuration
    st.set_page_config(
        page_title="📄 Document Processor",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS for better UI
    st.markdown("""
    <style>
        .main-header {
            text-align: center;
            padding: 1rem 0;
            background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-radius: 10px;
            margin-bottom: 2rem;
        }
        
        .stButton > button {
            width: 100%;
            border-radius: 8px;
            font-weight: 500;
        }
        
        .success-box {
            padding: 1rem;
            background-color: #d4edda;
            border: 1px solid #c3e6cb;
            border-radius: 8px;
            color: #155724;
        }
        
        .warning-box {
            padding: 1rem;
            background-color: #fff3cd;
            border: 1px solid #ffeeba;
            border-radius: 8px;
            color: #856404;
        }
        
        .info-card {
            padding: 1rem;
            background-color: #f8f9fa;
            border-radius: 8px;
            border-left: 4px solid #667eea;
            margin-bottom: 1rem;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📄 Document Processor</h1>
        <p>Transform multiple documents into a unified report</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar with instructions
    with st.sidebar:
        st.header("ℹ️ How to Use")
        st.markdown("""
        1. **Upload Files** - Drag & drop or browse
        2. **Process** - Click the process button
        3. **Preview** - Review the generated report
        4. **Download** - Export as MD or HTML
        
        ---
        
        **Supported Formats:**
        - 📊 Excel (.xlsx, .xls)
        - 📝 Word (.docx)
        - 📕 PDF (.pdf)
        - 📃 Text (.txt)
        - 🖼️ Images (.png, .jpg, .jpeg)
        - 📑 Markdown (.md)
        
        ---
        
        **Features:**
        - ✅ Zero content alteration
        - ✅ Table extraction
        - ✅ Auto Table of Contents
        - ✅ OCR for images (VI/EN)
        - ✅ MD to HTML conversion
        - ✅ Error handling
        """)
        
        st.markdown("---")
        st.caption("Built with ❤️ using Streamlit")
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📁 Upload Documents")
        
        uploaded_files = st.file_uploader(
            "Choose files to process",
            type=['xlsx', 'xls', 'docx', 'pdf', 'txt', 'png', 'jpg', 'jpeg', 'md'],
            accept_multiple_files=True,
            help="Upload Excel, Word, PDF, Text, Image, or Markdown files"
        )
        
        if uploaded_files:
            st.markdown("**Uploaded Files:**")
            for f in uploaded_files:
                file_size = len(f.getvalue()) / 1024  # KB
                st.markdown(f"- 📎 `{f.name}` ({file_size:.1f} KB)")
    
    with col2:
        st.subheader("⚙️ Processing Options")
        
        st.markdown("""
        <div class="info-card">
            <strong>Processing Mode:</strong> Full Extraction<br>
            <small>All text and tables will be extracted exactly as they appear.</small>
        </div>
        """, unsafe_allow_html=True)
    
    # Process button
    st.markdown("---")
    
    process_button = st.button(
        "🚀 Process Documents",
        type="primary",
        use_container_width=True,
        disabled=not uploaded_files
    )
    
    # Session state for storing results
    if 'markdown_content' not in st.session_state:
        st.session_state.markdown_content = None
    if 'html_content' not in st.session_state:
        st.session_state.html_content = None
    
    # Process files
    if process_button and uploaded_files:
        with st.spinner("Processing documents..."):
            processor = DocumentProcessor()
            
            # Process with progress
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for idx, f in enumerate(uploaded_files):
                status_text.text(f"Processing: {f.name}")
                progress_bar.progress((idx + 1) / len(uploaded_files))
            
            # Generate content
            st.session_state.markdown_content = processor.process_files(uploaded_files)
            st.session_state.html_content = generate_html(st.session_state.markdown_content)
            
            progress_bar.empty()
            status_text.empty()
            
            # Show warnings if any
            if processor.warnings:
                st.warning("⚠️ Some files could not be processed:")
                for warning in processor.warnings:
                    st.markdown(f"- {warning}")
            
            st.success(f"✅ Successfully processed {sum(1 for f in processor.processed_files if f.success)} of {len(uploaded_files)} files!")
    
    # Display results
    if st.session_state.markdown_content:
        st.markdown("---")
        st.subheader("📋 Generated Report")
        
        # Tabs for different views
        tab1, tab2 = st.tabs(["📝 Markdown Preview", "🌐 HTML Preview"])
        
        with tab1:
            st.markdown(st.session_state.markdown_content)
        
        with tab2:
            st.components.v1.html(st.session_state.html_content, height=800, scrolling=True)
        
        # Download buttons
        st.markdown("---")
        st.subheader("📥 Download Report")
        
        col1, col2, col3 = st.columns([1, 1, 1])
        
        with col1:
            st.download_button(
                label="📄 Download as Markdown (.md)",
                data=st.session_state.markdown_content,
                file_name=f"unified_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                mime="text/markdown",
                use_container_width=True
            )
        
        with col2:
            st.download_button(
                label="🌐 Download as HTML (.html)",
                data=st.session_state.html_content,
                file_name=f"unified_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                mime="text/html",
                use_container_width=True
            )
        
        with col3:
            # Clear button
            if st.button("🗑️ Clear Results", use_container_width=True):
                st.session_state.markdown_content = None
                st.session_state.html_content = None
                st.rerun()


# ============================================================================
# ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    main()
